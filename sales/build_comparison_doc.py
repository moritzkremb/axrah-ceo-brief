from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

RED = RGBColor(0xC8, 0x10, 0x2E)
NAVY = RGBColor(0x0F, 0x1B, 0x2D)
NAVY_MID = RGBColor(0x1A, 0x2B, 0x40)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x00, 0xC8, 0x78)
LIGHT_GREY = RGBColor(0xF4, 0xF6, 0xF8)
MID_GREY = RGBColor(0xE2, 0xE8, 0xEE)
DARK_GREY = RGBColor(0x44, 0x5A, 0x6A)
TEXT_MUTED = RGBColor(0x6B, 0x7C, 0x93)

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), val.get('val', 'single'))
            border.set(qn('w:sz'), val.get('sz', '4'))
            border.set(qn('w:color'), val.get('color', 'auto'))
            tcBorders.append(border)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, color=None, size=None, italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run

doc = Document()

# --- Page setup: landscape, narrow margins ---
section = doc.sections[0]
section.page_width = Inches(11)
section.page_height = Inches(8.5)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)
section.top_margin = Cm(1.2)
section.bottom_margin = Cm(1.2)

# Remove default paragraph spacing
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(0)

# ── HEADER ──────────────────────────────────────────────────────────────────
header_table = doc.add_table(rows=1, cols=2)
header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
header_table.style = 'Table Grid'
header_table.autofit = False
w = Inches(9.3)
header_table.columns[0].width = Inches(5.8)
header_table.columns[1].width = Inches(3.5)

lc = header_table.cell(0, 0)
set_cell_bg(lc, NAVY)
lp = lc.paragraphs[0]
lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(lp, 'AXRAH', bold=True, color=RED, size=22)
add_run(lp, '  vs. The Competition', bold=False, color=WHITE, size=14)
lp2 = lc.add_paragraph()
add_run(lp2, 'Full-body photobiomodulation — head-to-head comparison', color=RGBColor(0x8E, 0xA0, 0xB4), size=9)

rc = header_table.cell(0, 1)
set_cell_bg(rc, NAVY_MID)
rp = rc.paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(rp, 'Competitor prices are estimates. Use "approximately" or\n"estimated" when quoting to clients.',
        color=TEXT_MUTED, size=8, italic=True)

doc.add_paragraph()

# ── ADVANTAGE CALLOUTS ───────────────────────────────────────────────────────
adv_table = doc.add_table(rows=2, cols=3)
adv_table.alignment = WD_TABLE_ALIGNMENT.CENTER
adv_table.style = 'Table Grid'
adv_table.autofit = False
col_w = Inches(3.1)
for col in adv_table.columns:
    col.width = col_w

advantages = [
    ('LED DENSITY ADVANTAGE', '6.2x', 'Chamber Ultra: 43,200 LEDs vs MitoPOD 6,960'),
    ('IRRADIANCE ADVANTAGE', '4.3x', '129 mW/cm2 vs MitoPOD ~30 mW/cm2'),
    ('PRICE vs. NOVOTHOR', '~1/3', 'Chamber Ultra ~$27K vs NovoTHOR est. $65K+'),
]

for i, (label, value, sub) in enumerate(advantages):
    c = adv_table.cell(0, i)
    set_cell_bg(c, RGBColor(0x2A, 0x10, 0x16))
    p1 = c.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p1, label, bold=True, color=RED, size=8)

    c2 = adv_table.cell(1, i)
    set_cell_bg(c2, RGBColor(0x1E, 0x0B, 0x10))
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, value + '\n', bold=True, color=WHITE, size=22)
    add_run(p2, sub, color=TEXT_MUTED, size=8)

doc.add_paragraph()

# ── MAIN COMPARISON TABLE ────────────────────────────────────────────────────
COLS = ['', 'Chamber\n$16,999', 'Chamber Ultra\n$26,999', 'NovoTHOR\n~est. $65K+', 'TheraLight\n~est. $45–85K', 'Prism Light Pod\n~est. $35K+', 'MitoPOD\nConsumer']

rows_data = [
    # (section_header, data_rows)
    ('PERFORMANCE', [
        ('Total LEDs',         ['13,680',    '43,200 ✓',  'N/A',        'N/A',          '17,000',     '6,960']),
        ('Irradiance',         ['Optimised', '129 mW/cm²','N/A public', 'N/A public',   'N/A public', '~30 mW/cm²']),
        ('Wavelengths',        ['4\n633/850/940/1060nm', '5\n633/660/810/940/1060nm', '2\n660/850nm', '4\nmultiple', '3\n630/660/850nm', '2–3\nvaries']),
        ('Pulse Technology',   ['—',         '1–10,000 Hz','—',          'Limited',      '—',          '—']),
    ]),
    ('COVERAGE & FORM FACTOR', [
        ('Full-body coverage',      ['✓', '✓', '✓', '✓', '✓', 'Portable']),
        ('No repositioning',        ['✓', '✓', '✓', '✓', '✓', '✗']),
        ('Immersive chamber design',['✓', '✓', '✓', '✓', 'Partial', '✗']),
    ]),
    ('CONTROLS & OPERATION', [
        ('Industrial touchscreen',      ['✓', '✓', 'Basic', '✓', 'Basic', '✗']),
        ('Wireless remote operation',   ['✓', '✓', '✗', '✗', '✗', '✗']),
        ('Programmable sessions',       ['✓', '✓', 'Limited', '✓', 'Limited', '✗']),
        ('Multi-language UI',           ['EN/DE/FR/IT/ES', 'EN/DE/FR/IT/ES', 'EN only', 'EN only', 'EN only', 'EN only']),
    ]),
    ('COMMERCIAL', [
        ('Est. device price',       ['$16,999', '$26,999', '~$65K+ est.', '~$45–85K est.', '~$35K+ est.', 'Consumer']),
        ('B2B / clinic fit',        ['✓', '✓', '✓', '✓', '✓', '✗']),
        ('Published clinical evidence', ['PBM category', 'PBM category', '✓ Manufacturer', 'Some', 'Some', '✗']),
    ]),
]

n_cols = len(COLS)
total_rows = 1  # header
for section_name, rows in rows_data:
    total_rows += 1 + len(rows)

tbl = doc.add_table(rows=total_rows, cols=n_cols)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'
tbl.autofit = False

# Column widths
col_widths = [Inches(1.55), Inches(1.25), Inches(1.35), Inches(1.2), Inches(1.2), Inches(1.2), Inches(1.2)]
for i, w in enumerate(col_widths):
    for cell in tbl.columns[i].cells:
        cell.width = w

# Header row
hrow = tbl.rows[0]
for ci, col_text in enumerate(COLS):
    cell = hrow.cells[ci]
    if ci == 0:
        set_cell_bg(cell, NAVY)
    elif ci in (1, 2):
        set_cell_bg(cell, RGBColor(0x2A, 0x10, 0x16))
    else:
        set_cell_bg(cell, NAVY_MID)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines = col_text.split('\n')
    add_run(p, lines[0], bold=True, color=WHITE if ci not in (1,2) else RED, size=9)
    if len(lines) > 1:
        add_run(p, '\n' + lines[1], bold=False, color=RGBColor(0x8E,0xA0,0xB4), size=8)

# Data rows
current_row = 1
for section_name, rows in rows_data:
    # Section header row
    srow = tbl.rows[current_row]
    merged = srow.cells[0].merge(srow.cells[n_cols - 1])
    set_cell_bg(merged, NAVY_MID)
    sp = merged.paragraphs[0]
    sp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(sp, section_name, bold=True, color=RGBColor(0x8E,0xA0,0xB4), size=8)
    current_row += 1

    for row_label, values in rows:
        drow = tbl.rows[current_row]
        # Row label
        lc = drow.cells[0]
        set_cell_bg(lc, LIGHT_GREY)
        lp = lc.paragraphs[0]
        add_run(lp, row_label, bold=True, color=DARK_GREY, size=8)

        for vi, val in enumerate(values):
            ci = vi + 1
            cell = drow.cells[ci]

            # Background
            if ci == 1:
                set_cell_bg(cell, RGBColor(0xFF, 0xF5, 0xF6))
            elif ci == 2:
                set_cell_bg(cell, RGBColor(0xFF, 0xEE, 0xEF))
            else:
                set_cell_bg(cell, WHITE)

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Colour logic
            lines = val.split('\n')
            main = lines[0]
            sub = lines[1] if len(lines) > 1 else None

            if main == '✓':
                add_run(p, main, bold=True, color=GREEN, size=10)
            elif main == '✗' or main == '—':
                add_run(p, main, bold=False, color=TEXT_MUTED, size=10)
            elif '✓' in main and ci == 2:  # Ultra winner field
                add_run(p, main.replace(' ✓',''), bold=True, color=GREEN, size=9)
            elif ci in (1, 2) and main not in ('—', 'N/A', 'Optimised', 'PBM category'):
                add_run(p, main, bold=True, color=RED if ci == 1 else RED, size=9)
            else:
                add_run(p, main, bold=False, color=DARK_GREY, size=8)

            if sub:
                add_run(p, '\n' + sub, bold=False, color=TEXT_MUTED, size=7)

        current_row += 1

doc.add_paragraph()

# ── BOTTOM CALLOUTS ──────────────────────────────────────────────────────────
bot_table = doc.add_table(rows=1, cols=2)
bot_table.alignment = WD_TABLE_ALIGNMENT.CENTER
bot_table.style = 'Table Grid'
bot_table.autofit = False
bot_table.columns[0].width = Inches(4.6)
bot_table.columns[1].width = Inches(4.6)

left_cell = bot_table.cell(0, 0)
right_cell = bot_table.cell(0, 1)
set_cell_bg(left_cell, LIGHT_GREY)
set_cell_bg(right_cell, LIGHT_GREY)

lp = left_cell.paragraphs[0]
add_run(lp, 'AXRAH CHAMBER ULTRA — KEY ADVANTAGES\n', bold=True, color=RED, size=9)
advantages_text = [
    ('43,200 LEDs', ' — 2.5× Prism Light Pod, 6.2× MitoPOD'),
    ('5 wavelengths', ' — broadest in class (633/660/810/940/1060nm)'),
    ('129 mW/cm²', ' — 4.3× the irradiance of MitoPOD'),
    ('Pulse 1–10,000 Hz', ' — no competitor at this price point'),
    ('~1/3 price of NovoTHOR', ' — for superior specification'),
    ('5-language UI', ' — built for global deployment'),
]
for bold_part, rest in advantages_text:
    add_run(lp, '• ', color=RED, size=8)
    add_run(lp, bold_part, bold=True, color=DARK_GREY, size=8)
    add_run(lp, rest + '\n', color=DARK_GREY, size=8)

rp = right_cell.paragraphs[0]
add_run(rp, 'B2B TARGET USE CASES\n', bold=True, color=RED, size=9)
use_cases = [
    ('Chamber ($16,999)', ' — Medspas, chiropractic, PT clinics, boutique gyms'),
    ('Chamber Ultra ($26,999)', ' — Performance centers, sports teams, luxury hotels,\n   neurological rehab, integrative medicine'),
    ('Session pricing', ' — $50–$100/session typical; payback in 6–18 months'),
    ('Wireless + programmable', ' — Built for staff-run clinical environments'),
]
for bold_part, rest in use_cases:
    add_run(rp, '• ', color=RED, size=8)
    add_run(rp, bold_part, bold=True, color=DARK_GREY, size=8)
    add_run(rp, rest + '\n', color=DARK_GREY, size=8)

# ── FOOTNOTE ─────────────────────────────────────────────────────────────────
doc.add_paragraph()
fp = doc.add_paragraph()
add_run(fp, 'Competitor prices are estimates based on publicly available sources and may vary by configuration, region, or distributor. Always present competitor figures as approximate. AXRAH device specs per official product pages (May 2026). Competitor LED counts and irradiance reflect publicly available data; "N/A public" indicates specifications not published by the manufacturer. © AXRAH 2026 — Internal use only.',
        color=TEXT_MUTED, size=7, italic=True)

out_path = '/home/user/axrah-ceo-brief/sales/AXRAH-Comparison-OnePager.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
