from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── PALETTE ───────────────────────────────────────────────────────────────────
RED      = RGBColor(0xC8, 0x10, 0x2E)
NAVY     = RGBColor(0x0F, 0x1B, 0x2D)
NAVY_MID = RGBColor(0x1A, 0x2B, 0x40)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
MUTED    = RGBColor(0x6B, 0x7C, 0x93)
DARK     = RGBColor(0x22, 0x34, 0x4A)
GREEN    = RGBColor(0x00, 0x7A, 0x4C)
AMBER    = RGBColor(0xB4, 0x5A, 0x00)

def hex_color(rgb: RGBColor):
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color(rgb))
    tcPr.append(shd)

def set_para_shading(para, rgb: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color(rgb))
    pPr.append(shd)

def add_run(para, text, bold=False, color=None, size=None, italic=False, underline=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if color: run.font.color.rgb = color
    if size:  run.font.size = Pt(size)
    run.font.name = 'Calibri'
    return run

doc = Document()

# ── PAGE SETUP ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Cm(2.0)
section.right_margin  = Cm(2.0)
section.top_margin    = Cm(1.8)
section.bottom_margin = Cm(1.8)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after  = Pt(0)

# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def spacer(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

def page_break():
    doc.add_page_break()

def heading_block(text, sub=None):
    """Full-width navy heading band."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, NAVY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, text, bold=True, color=WHITE, size=15)
    if sub:
        add_run(p, f'\n{sub}', bold=False, color=MUTED, size=9)
    spacer(1)

def section_label(text, color=RED):
    """Inline section tag — e.g. MEDSPA / AESTHETIC CLINIC"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, f'  {text}  ', bold=True, color=WHITE, size=9)
    spacer(1)

def tag(para, text, bg_color=RED):
    """Inline coloured pill tag using a 1-cell table hack via run background."""
    add_run(para, f' {text} ', bold=True, color=WHITE, size=8)

def script_box(lines, box_color=RGBColor(0xF0, 0xF6, 0xFF),
               border_color=RGBColor(0x1A, 0x2B, 0x40), label=None):
    """Shaded box for script content."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, box_color)

    first = True
    for line in lines:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.15)

        if isinstance(line, tuple):
            kind, text = line
            if kind == 'label':
                add_run(p, text, bold=True, color=DARK, size=8)
            elif kind == 'body':
                add_run(p, text, color=DARK, size=10)
            elif kind == 'italic':
                add_run(p, text, italic=True, color=MUTED, size=9)
            elif kind == 'bold':
                add_run(p, text, bold=True, color=DARK, size=10)
            elif kind == 'blank':
                add_run(p, '', size=6)
        else:
            add_run(p, line, color=DARK, size=10)

    spacer(1)

def two_col(left_label, left_lines, right_label, right_lines):
    """Two-column script layout."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    tbl.autofit = False
    tbl.columns[0].width = Inches(3.6)
    tbl.columns[1].width = Inches(3.6)

    for ci, (lbl, lines) in enumerate([(left_label, left_lines), (right_label, right_lines)]):
        cell = tbl.cell(0, ci)
        bg = RGBColor(0xF0, 0xF6, 0xFF) if ci == 0 else RGBColor(0xF0, 0xFB, 0xF5)
        set_cell_bg(cell, bg)

        # Label row
        p = cell.paragraphs[0]
        p.paragraph_format.left_indent = Inches(0.12)
        lbl_color = NAVY_MID if ci == 0 else GREEN
        add_run(p, lbl, bold=True, color=lbl_color, size=8)

        for line in lines:
            lp = cell.add_paragraph()
            lp.paragraph_format.left_indent  = Inches(0.12)
            lp.paragraph_format.right_indent = Inches(0.12)
            if isinstance(line, tuple):
                kind, text = line
                if kind == 'blank': add_run(lp, '', size=5)
                elif kind == 'italic': add_run(lp, text, italic=True, color=MUTED, size=9)
                elif kind == 'bold': add_run(lp, text, bold=True, color=DARK, size=10)
                else: add_run(lp, text, color=DARK, size=10)
            else:
                add_run(lp, line, color=DARK, size=10)

    spacer(1)

def inline_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    add_run(p, '→  ', bold=True, color=RED, size=9)
    add_run(p, text, italic=True, color=MUTED, size=9)
    spacer(1)

def objection_row(obj, response):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    tbl.autofit = False
    tbl.columns[0].width = Inches(2.5)
    tbl.columns[1].width = Inches(4.7)

    lc = tbl.cell(0, 0)
    set_cell_bg(lc, RGBColor(0xFF, 0xF0, 0xF2))
    lp = lc.paragraphs[0]
    lp.paragraph_format.left_indent = Inches(0.1)
    add_run(lp, obj, bold=True, color=RED, size=9)

    rc = tbl.cell(0, 1)
    set_cell_bg(rc, RGBColor(0xF7, 0xF9, 0xFC))
    rp = rc.paragraphs[0]
    rp.paragraph_format.left_indent  = Inches(0.1)
    rp.paragraph_format.right_indent = Inches(0.1)
    add_run(rp, response, color=DARK, size=9.5)

    spacer(1)

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

spacer(3)

tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'
cell = tbl.cell(0, 0)
set_cell_bg(cell, NAVY)

p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'AXRAH', bold=True, color=RED, size=42)
add_run(p, '\nSALES BIBLE', bold=True, color=WHITE, size=24)
add_run(p, '\n\nInitial Outreach Scripts — Email & SMS', bold=False, color=MUTED, size=12)
add_run(p, '\n\nFor internal use only  ·  2026', bold=False, color=MUTED, size=9)

spacer(2)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p2, 'Chamber  $16,999   ·   Chamber Ultra  $26,999', bold=True, color=DARK, size=11)

spacer(1)
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p3,
    'Medspas  ·  Gyms  ·  Chiro / PT  ·  Sports Teams  ·  Hotels  ·  Integrative Clinics',
    color=MUTED, size=9)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# HOW TO USE
# ══════════════════════════════════════════════════════════════════════════════

heading_block('HOW TO USE THIS GUIDE',
              'Read this before you send a single message.')

rules = [
    ('1.', 'Personalise the first line.',
     'Every script has a [bracket]. That bracket is your job. '
     'Generic first lines get deleted. A specific observation about their business gets read.'),
    ('2.', 'One ask per message.',
     'The only CTA is a 10–15 minute call. Not a demo, not a proposal, not a decision. '
     'Lower the bar and conversion goes up.'),
    ('3.', 'Send the email first, text 48–72 hours later.',
     'Text messages cold are intrusive unless they reference something already sent. '
     'The email gives the text legitimacy.'),
    ('4.', 'Do not pitch the product.',
     'You are selling a conversation, not a Chamber. '
     'The product spec comes on the call. In the outreach, you are selling the business outcome.'),
    ('5.', 'Use the ROI model before you call.',
     'Run their numbers in the Excel calculator before outreach. '
     'When you say "at 6 sessions a day you\'re looking at payback in 8 months" and you\'re right, '
     'you sound like a consultant, not a salesperson.'),
    ('6.', 'Follow up exactly once at Day 3.',
     'The Day-3 follow-up is included for every vertical. After that, move on. '
     'Chasing more than twice signals desperation and damages the brand.'),
]

for num, title, body in rules:
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    tbl.autofit = False
    tbl.columns[0].width = Inches(0.35)
    tbl.columns[1].width = Inches(6.85)

    nc = tbl.cell(0, 0)
    set_cell_bg(nc, RED)
    np_ = nc.paragraphs[0]
    np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(np_, num, bold=True, color=WHITE, size=10)

    bc = tbl.cell(0, 1)
    set_cell_bg(bc, RGBColor(0xF7, 0xF9, 0xFC))
    bp = bc.paragraphs[0]
    bp.paragraph_format.left_indent = Inches(0.1)
    add_run(bp, title + '  ', bold=True, color=DARK, size=10)
    add_run(bp, body, color=DARK, size=10)

    spacer(1)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# VERTICAL SCRIPTS — shared structure
# ══════════════════════════════════════════════════════════════════════════════

VERTICALS = [

  {
    'name': 'MEDSPA / AESTHETIC CLINIC',
    'color': RED,
    'context': (
        'Pain point: they have empty room hours and no passive revenue stream. '
        'Lens: ROI and revenue per square foot. '
        'Contact: Owner, Practice Manager, or Spa Director.'
    ),
    'subject_a': 'Adding a revenue line to [Clinic Name]',
    'subject_b': 'Quick question about [Clinic Name]\'s treatment menu',
    'email': [
        ('blank', ''),
        ('body', 'Hi [First Name],'),
        ('blank', ''),
        ('body',
         'Most medspas in [City] are sitting on an untapped revenue line: '
         'full-body red light therapy sessions at $75–100 each, '
         '15 minutes, zero staff interaction.'),
        ('blank', ''),
        ('body',
         'The AXRAH Chamber installs in any existing treatment room, '
         'runs autonomously, and pays for itself in under 9 months at moderate utilisation. '
         'Our clients typically add $6–10K/month in new revenue with no new hires.'),
        ('blank', ''),
        ('body',
         'Would it be useful if I pulled together a quick revenue model '
         'for [Clinic Name] specifically?'),
        ('blank', ''),
        ('body', '[Your Name]\nAXRAH'),
        ('blank', ''),
    ],
    'sms': (
        'Hi [First Name], [Your Name] from AXRAH. Quick one — is [Clinic Name] '
        'offering full-body red light therapy? Our medspa clients are adding $6–10K/month '
        'with zero extra staff. Worth a 10-min call?'
    ),
    'followup': (
        'Hi [First Name] — just bumping this up in case it got buried. '
        'I put together a quick revenue model for a medspa your size in [City] — '
        'happy to share, takes 5 minutes. Still the right person to speak to?'
    ),
    'personalise': [
        'Mention a specific treatment they already offer (e.g. "I see you do CoolSculpting — clients who do body treatments are the exact profile who book red light.")',
        'Reference their location if you know the competitive set ("Three medspas in [neighbourhood] have added RLT in the last 6 months.")',
        'Use their Google review count as a proxy for volume ("With [X] reviews, you\'re clearly doing strong session volume already.")',
    ],
  },

  {
    'name': 'GYM / PERFORMANCE CENTER',
    'color': NAVY_MID,
    'context': (
        'Pain point: losing recovery revenue to medspas. '
        'Lens: new revenue without new staff or square footage changes. '
        'Contact: Owner, General Manager, or Head of Performance.'
    ),
    'subject_a': 'The recovery revenue [Gym Name] isn\'t capturing',
    'subject_b': 'Your members are paying your competitor for recovery',
    'email': [
        ('blank', ''),
        ('body', 'Hi [First Name],'),
        ('blank', ''),
        ('body',
         'Recovery is the fastest-growing upsell in performance fitness — '
         'and most gym owners watch their members drive to a medspa for it '
         'instead of spending that money in-house.'),
        ('blank', ''),
        ('body',
         'The AXRAH Chamber installs in any recovery room, runs autonomously, '
         'and generates $50–100 per 15-minute session. '
         'No new hires. Most operators hit break-even in under 12 months.'),
        ('blank', ''),
        ('body',
         'Can I send you a quick revenue model built around '
         '[Gym Name]\'s session capacity?'),
        ('blank', ''),
        ('body', '[Your Name]\nAXRAH'),
        ('blank', ''),
    ],
    'sms': (
        'Hi [First Name], [Your Name] at AXRAH. Are [Gym Name] members asking about '
        'red light therapy? We help gyms capture $5–8K/month in recovery revenue '
        'they\'re currently sending elsewhere. 10-min call?'
    ),
    'followup': (
        'Hi [First Name] — following up from a few days ago. '
        'I built a quick revenue model for a gym your size — '
        'happy to walk you through it in under 10 minutes. Worth a look?'
    ),
    'personalise': [
        'Reference their class types ("Seeing you run a lot of HIIT and weightlifting — those members are your highest-intent recovery buyers.")',
        'Mention their membership size if visible ("A gym with [X] members typically converts 8–12% to recovery sessions.")',
        'Acknowledge existing recovery offerings ("If you already have cold plunge or sauna, RLT completes the stack — clients use all three.")',
    ],
  },

  {
    'name': 'CHIROPRACTIC / PHYSICAL THERAPY',
    'color': RGBColor(0x0A, 0x52, 0x78),
    'context': (
        'Pain point: aware of clinical PBM but priced out by NovoTHOR/TheraLight. '
        'Lens: clinical credibility + ROI within practice. '
        'Contact: Practice Owner or Clinical Director.'
    ),
    'subject_a': 'Clinical-grade PBM at 1/3 the price of NovoTHOR',
    'subject_b': 'Full-body photobiomodulation — without the $65K price tag',
    'email': [
        ('blank', ''),
        ('body', 'Hi [First Name],'),
        ('blank', ''),
        ('body',
         'If you\'ve looked at NovoTHOR or TheraLight and stepped back because of the price, '
         'the AXRAH Chamber delivers clinical-grade whole-body PBM at approximately $17K — '
         'with five wavelengths versus NovoTHOR\'s two.'),
        ('blank', ''),
        ('body',
         'Most of our chiro and PT clients add $4–8K/month in session revenue '
         'and recover the device cost within the first year. '
         'Recent RCT data on PBM and cognitive recovery has also opened new patient conversations.'),
        ('blank', ''),
        ('body',
         'Happy to send the spec comparison and a revenue model for a practice your size — '
         'would that be useful?'),
        ('blank', ''),
        ('body', '[Your Name]\nAXRAH'),
        ('blank', ''),
    ],
    'sms': (
        'Hi [First Name], [Your Name] from AXRAH. Have you looked at full-body PBM for '
        '[Practice Name]? Clinical-grade, 5 wavelengths, ~$17K vs NovoTHOR\'s $65K+. '
        'Most practices recover cost in <12 months. Worth a quick call?'
    ),
    'followup': (
        'Hi [First Name] — following up on my note about full-body PBM. '
        'I put together a side-by-side of AXRAH vs NovoTHOR and a simple revenue model. '
        'Happy to share — still relevant for [Practice Name]?'
    ),
    'personalise': [
        'Reference their specific speciality ("Seeing you focus on sports rehab — the muscle recovery and cognitive performance data on PBM is directly relevant for your athlete patients.")',
        'Mention CPT billing if relevant to their market ("Depending on your payer mix, some practices are billing PBM sessions — happy to share what we\'ve seen.")',
        'Reference their equipment ("Practitioners with laser/shockwave therapy typically see the strongest patient uptake for PBM as a complement.")',
    ],
  },

  {
    'name': 'SPORTS TEAM — ATHLETIC / PERFORMANCE DIRECTOR',
    'color': RGBColor(0x00, 0x52, 0x1F),
    'context': (
        'Pain point: recovery throughput and edge in cognitive performance. '
        'Lens: competitive advantage + ROI vs NovoTHOR budget. '
        'Contact: Head of Athletic Training, Performance Director, or GM.'
    ),
    'subject_a': 'Recovery + cognitive performance for [Team Name]',
    'subject_b': 'What the top performance programs are adding to their recovery stack',
    'email': [
        ('blank', ''),
        ('body', 'Hi [First Name],'),
        ('blank', ''),
        ('body',
         'New clinical data shows whole-body near-infrared PBM significantly improves '
         'cognitive scores over 12 weeks — on top of the established muscle recovery '
         'and inflammation reduction benefits. It\'s becoming a serious tool '
         'for performance programs that want an edge beyond standard recovery protocols.'),
        ('blank', ''),
        ('body',
         'The AXRAH Chamber Ultra gives athletes full-body treatment across five wavelengths '
         'in a single 15-minute session. It\'s what performance centers are choosing '
         'instead of NovoTHOR — same clinical-grade output, approximately 1/3 the price.'),
        ('blank', ''),
        ('body',
         'Would a 15-minute call with your performance staff be worth arranging?'),
        ('blank', ''),
        ('body', '[Your Name]\nAXRAH'),
        ('blank', ''),
    ],
    'sms': (
        'Hi [First Name], [Your Name] from AXRAH. We work with performance programs on '
        'whole-body red light therapy — recovery + cognitive performance in 15 min/session. '
        'Clinical-grade, used instead of NovoTHOR at 1/3 the cost. Worth a quick call?'
    ),
    'followup': (
        'Hi [First Name] — checking back in. '
        'Happy to share the cognitive performance clinical data and a spec sheet '
        'built for sports performance environments. Still the right contact for this at [Team Name]?'
    ),
    'personalise': [
        'Reference their sport and season ("Coming into pre-season camp, recovery throughput is typically the constraint — the Chamber processes a full squad in under 2 hours.")',
        'Name a competitor or peer program ("Several [league] programs have installed whole-body PBM in the last 18 months — happy to share which if useful.")',
        'Acknowledge existing setup ("If you already have cold plunge and compression, PBM is the piece that addresses the cellular recovery layer those tools miss.")',
    ],
  },

  {
    'name': 'HOTEL / RESORT SPA',
    'color': RGBColor(0x5A, 0x3E, 0x00),
    'context': (
        'Pain point: undifferentiated spa menu, pressure to add premium amenities. '
        'Lens: signature treatment + high revenue per minute. '
        'Contact: Spa Director, Director of Wellness, or F&B/Amenities GM.'
    ),
    'subject_a': 'A signature $100/session add-on for [Hotel Name]\'s spa',
    'subject_b': 'What [Hotel Name]\'s spa competitors are already offering',
    'email': [
        ('blank', ''),
        ('body', 'Hi [First Name],'),
        ('blank', ''),
        ('body',
         'Full-body red light therapy is becoming the premium wellness amenity that '
         '4-star+ spa directors are adding to their menus — '
         '$75–120 per 15-minute session, no consumables, no therapist required.'),
        ('blank', ''),
        ('body',
         'The AXRAH Chamber Ultra is clinical-grade and visually striking — '
         'it positions naturally as a signature treatment rather than a commodity service. '
         'Five wavelengths, programmable protocols, multi-language touchscreen for international guests.'),
        ('blank', ''),
        ('body',
         'Could we arrange a 15-minute call to walk through how it fits [Hotel Name]\'s wellness offering?'),
        ('blank', ''),
        ('body', '[Your Name]\nAXRAH'),
        ('blank', ''),
    ],
    'sms': (
        'Hi [First Name], [Your Name] from AXRAH. Is [Hotel Name]\'s spa considering '
        'full-body red light therapy? We\'re placing clinical-grade Chambers with luxury '
        'hotel spas at $75–120/session. Happy to share details — worth 15 min?'
    ),
    'followup': (
        'Hi [First Name] — following up on the red light therapy note. '
        'Happy to send a one-page overview of how other luxury properties are '
        'positioning it as a signature treatment. Useful?'
    ),
    'personalise': [
        'Reference their star rating or brand positioning ("Given [Hotel Name]\'s 5-star positioning, the Chamber Ultra\'s design and spec level are the right fit — the standard Chamber is better suited to a more utilitarian wellness space.")',
        'Mention their existing spa menu ("I see you offer [treatment] — red light complements that well as a pre- or post-treatment protocol.")',
        'Reference guest demographics ("A resort with your international guest mix benefits from the multi-language interface — EN/DE/FR/IT/ES out of the box.")',
    ],
  },

  {
    'name': 'INTEGRATIVE / LONGEVITY CLINIC',
    'color': RGBColor(0x2A, 0x5C, 0x45),
    'context': (
        'Pain point: clients asking about PBM stacking; gap in protocol menu. '
        'Lens: premium protocol bundle + mitochondrial science. '
        'Contact: Medical Director, Clinic Owner, or Head of Protocols.'
    ),
    'subject_a': 'Stacking PBM with your existing protocols — the revenue case',
    'subject_b': 'Full-body red light + your IV and peptide protocols',
    'email': [
        ('blank', ''),
        ('body', 'Hi [First Name],'),
        ('blank', ''),
        ('body',
         'Your clients on methylene blue or NAD+ IVs are increasingly asking about '
         'stacking with photobiomodulation — the data on combined mitochondrial activation '
         'is compelling, and the protocol is starting to become a differentiator for '
         'the clinics that offer it.'),
        ('blank', ''),
        ('body',
         'The AXRAH Chamber Ultra is what integrative clinics are installing to offer '
         'a premium PBM + MB protocol at $150–200 per combined session. '
         'Five wavelengths, programmable pulse frequencies, clinical-grade irradiance.'),
        ('blank', ''),
        ('body',
         'Would a quick call make sense to explore how it fits your current protocol menu?'),
        ('blank', ''),
        ('body', '[Your Name]\nAXRAH'),
        ('blank', ''),
    ],
    'sms': (
        'Hi [First Name], [Your Name] from AXRAH. Are your clients asking about '
        'red light stacking with their IV protocols? We work with integrative clinics '
        'on PBM + MB stacks at $150–200/session. Worth a quick call?'
    ),
    'followup': (
        'Hi [First Name] — just following up. '
        'Happy to share a protocol guide and revenue model for the PBM + MB stack specifically. '
        'Still on your radar for [Clinic Name]?'
    ),
    'personalise': [
        'Reference specific protocols they advertise ("Seeing you offer NAD+ IV — that\'s exactly the client profile who asks about PBM. The mitochondrial activation overlap is the key conversation.")',
        'Name the science angle ("The methylene blue + PBM combination is getting serious attention in the longevity community right now — Bryan Johnson\'s protocol circle is talking about it.")',
        'Reference their pricing ("At your IV pricing, a $150–200 PBM stack session fits your client spend profile without friction.")',
    ],
  },

]

# ══════════════════════════════════════════════════════════════════════════════
# RENDER EACH VERTICAL
# ══════════════════════════════════════════════════════════════════════════════

for vi, v in enumerate(VERTICALS):

    # Section banner
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, v['color'])
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, f'  {vi+1} / 6  ·  {v["name"]}  ', bold=True, color=WHITE, size=12)
    spacer(1)

    # Context note
    ctx_tbl = doc.add_table(rows=1, cols=1)
    ctx_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    ctx_tbl.style = 'Table Grid'
    ctx_cell = ctx_tbl.cell(0, 0)
    set_cell_bg(ctx_cell, RGBColor(0xF7, 0xF9, 0xFC))
    cp = ctx_cell.paragraphs[0]
    cp.paragraph_format.left_indent = Inches(0.12)
    add_run(cp, v['context'], italic=True, color=MUTED, size=9)
    spacer(1)

    # ── EMAIL ─────────────────────────────────────────────────────────────────

    # Subject lines
    subj_tbl = doc.add_table(rows=2, cols=2)
    subj_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    subj_tbl.style = 'Table Grid'
    subj_tbl.autofit = False
    subj_tbl.columns[0].width = Inches(1.1)
    subj_tbl.columns[1].width = Inches(6.1)

    for ri, (tag_text, subj_text) in enumerate([
        ('SUBJECT A', v['subject_a']),
        ('SUBJECT B', v['subject_b']),
    ]):
        lc = subj_tbl.cell(ri, 0)
        set_cell_bg(lc, RGBColor(0xE8, 0xEC, 0xF2))
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(lp, tag_text, bold=True, color=NAVY_MID, size=8)

        rc = subj_tbl.cell(ri, 1)
        set_cell_bg(rc, WHITE)
        rp = rc.paragraphs[0]
        rp.paragraph_format.left_indent = Inches(0.1)
        add_run(rp, subj_text, bold=True, color=DARK, size=10)

    spacer(1)

    # Email body
    email_tbl = doc.add_table(rows=1, cols=1)
    email_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    email_tbl.style = 'Table Grid'
    email_cell = email_tbl.cell(0, 0)
    set_cell_bg(email_cell, RGBColor(0xF0, 0xF6, 0xFF))

    header_p = email_cell.paragraphs[0]
    header_p.paragraph_format.left_indent = Inches(0.15)
    add_run(header_p, 'EMAIL — COLD OUTREACH', bold=True, color=NAVY_MID, size=8)

    for line in v['email']:
        lp = email_cell.add_paragraph()
        lp.paragraph_format.left_indent  = Inches(0.15)
        lp.paragraph_format.right_indent = Inches(0.15)
        kind, text = line
        if kind == 'blank':
            lp.paragraph_format.space_before = Pt(2)
            lp.paragraph_format.space_after  = Pt(2)
        elif kind == 'bold':
            add_run(lp, text, bold=True, color=DARK, size=10)
        elif kind == 'italic':
            add_run(lp, text, italic=True, color=MUTED, size=9)
        else:
            add_run(lp, text, color=DARK, size=10)

    spacer(1)

    # ── SMS ───────────────────────────────────────────────────────────────────
    sms_tbl = doc.add_table(rows=1, cols=1)
    sms_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sms_tbl.style = 'Table Grid'
    sms_cell = sms_tbl.cell(0, 0)
    set_cell_bg(sms_cell, RGBColor(0xF0, 0xFB, 0xF5))

    sp1 = sms_cell.paragraphs[0]
    sp1.paragraph_format.left_indent = Inches(0.15)
    add_run(sp1, 'TEXT / SMS — send 48–72 hrs after email', bold=True, color=GREEN, size=8)

    sp2 = sms_cell.add_paragraph()
    sp2.paragraph_format.left_indent  = Inches(0.15)
    sp2.paragraph_format.right_indent = Inches(0.15)
    add_run(sp2, v['sms'], color=DARK, size=10)

    spacer(1)

    # ── DAY-3 FOLLOW-UP ───────────────────────────────────────────────────────
    fu_tbl = doc.add_table(rows=1, cols=1)
    fu_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    fu_tbl.style = 'Table Grid'
    fu_cell = fu_tbl.cell(0, 0)
    set_cell_bg(fu_cell, RGBColor(0xFB, 0xF7, 0xFF))

    fp1 = fu_cell.paragraphs[0]
    fp1.paragraph_format.left_indent = Inches(0.15)
    add_run(fp1, 'DAY-3 EMAIL FOLLOW-UP — reply to the original thread', bold=True, color=RGBColor(0x5A, 0x20, 0x8A), size=8)

    fp2 = fu_cell.add_paragraph()
    fp2.paragraph_format.left_indent  = Inches(0.15)
    fp2.paragraph_format.right_indent = Inches(0.15)
    add_run(fp2, v['followup'], color=DARK, size=10)

    spacer(1)

    # ── PERSONALISATION NOTES ────────────────────────────────────────────────
    pers_tbl = doc.add_table(rows=1, cols=1)
    pers_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    pers_tbl.style = 'Table Grid'
    pers_cell = pers_tbl.cell(0, 0)
    set_cell_bg(pers_cell, RGBColor(0xFB, 0xF5, 0xEC))

    pp1 = pers_cell.paragraphs[0]
    pp1.paragraph_format.left_indent = Inches(0.15)
    add_run(pp1, 'HOW TO PERSONALISE', bold=True, color=RGBColor(0x8A, 0x50, 0x00), size=8)

    for note in v['personalise']:
        np_ = pers_cell.add_paragraph()
        np_.paragraph_format.left_indent  = Inches(0.25)
        np_.paragraph_format.right_indent = Inches(0.15)
        add_run(np_, '•  ', bold=True, color=RED, size=9)
        add_run(np_, note, color=DARK, size=9)

    spacer(2)

    if vi < len(VERTICALS) - 1:
        page_break()

# ══════════════════════════════════════════════════════════════════════════════
# OBJECTION HANDLING
# ══════════════════════════════════════════════════════════════════════════════

page_break()

heading_block('OBJECTION HANDLING',
              'These are for the first call, not the email. Have them ready.')

OBJECTIONS = [
    (
        '"We\'re not interested."',
        'Completely understand. Would it be okay if I sent a one-page overview for when the timing changes? These conversations tend to come back around — and when they do, it\'s usually after a competitor has installed one.'
    ),
    (
        '"We already have a red light panel / device."',
        'Good to know — actually most of our clients who came from panels saw a 3–4x jump in utilisation and revenue when they moved to a full-body chamber. The session experience is different enough that clients treat it as a separate service. Happy to share a quick comparison if useful?'
    ),
    (
        '"It\'s too expensive / not in the budget."',
        'I hear that a lot — the number that usually shifts the conversation is the payback period. At [X sessions/week] and $[Y]/session, you\'re typically looking at break-even in [Z] months. Want me to run those exact numbers for your volume before you decide?'
    ),
    (
        '"We need to think about it / not ready yet."',
        'Of course. What would need to be true for this to be a priority in the next quarter? I want to make sure I come back at the right time with the right information — not just to check in.'
    ),
    (
        '"Send me some information."',
        'Absolutely — before I do, quick question: what\'s your current session capacity and what are you typically charging per treatment? I\'ll tailor the revenue model to your numbers so it\'s actually useful rather than generic.'
    ),
    (
        '"We already have a NovoTHOR."',
        'That\'s great — you\'re clearly already committed to clinical PBM at the highest level. Two questions: are you running at full utilisation? And do you have a second location or expansion planned? Some of our NovoTHOR clients add a Chamber to handle overflow demand or as a lower-price-point entry for a new site.'
    ),
    (
        '"I\'ve never heard of AXRAH."',
        'That\'s fair — we\'re a newer entrant in a market that\'s been dominated by $65K devices. The reason we\'re getting attention is the spec: 43,200 LEDs, five wavelengths, 129 mW/cm² — numbers that beat the established players — at a price point that makes the ROI straightforward. Happy to send the comparison sheet.'
    ),
]

for obj, resp in OBJECTIONS:
    objection_row(obj, resp)

# ── CADENCE SUMMARY ────────────────────────────────────────────────────────
spacer(1)
heading_block('OUTREACH CADENCE', 'The full sequence in one view.')

cad_tbl = doc.add_table(rows=5, cols=3)
cad_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cad_tbl.style = 'Table Grid'
cad_tbl.autofit = False
cad_tbl.columns[0].width = Inches(1.2)
cad_tbl.columns[1].width = Inches(2.2)
cad_tbl.columns[2].width = Inches(3.8)

for ci, hdr in enumerate(['Timing', 'Action', 'Goal']):
    cell = cad_tbl.cell(0, ci)
    set_cell_bg(cell, NAVY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, hdr, bold=True, color=WHITE, size=9)

cadence = [
    ('Day 1',      'Cold email',        'Get on their radar. Plant the revenue idea.'),
    ('Day 3',      'Follow-up email',   'One gentle bump. Reply to original thread.'),
    ('Day 3–4',    'SMS',               'Reference the email. Keep it human and brief.'),
    ('Day 10–14',  'Call (if no reply)', 'Reference the emails. Ask one qualifying question.'),
]

for ri, (timing, action, goal) in enumerate(cadence, 1):
    bg = RGBColor(0xFF, 0xF0, 0xF2) if ri % 2 == 0 else RGBColor(0xF7, 0xF9, 0xFC)
    for ci, text in enumerate([timing, action, goal]):
        cell = cad_tbl.cell(ri, ci)
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.left_indent = Inches(0.08)
        add_run(p, text, bold=(ci == 0), color=DARK if ci < 2 else NAVY_MID, size=9)

spacer(2)

# Footer
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(fp, 'AXRAH Sales Bible  ·  Initial Outreach Scripts  ·  Internal use only  ·  2026',
        italic=True, color=MUTED, size=8)

# ── SAVE ──────────────────────────────────────────────────────────────────────
out = '/home/user/axrah-ceo-brief/sales/AXRAH-Sales-Bible.docx'
doc.save(out)
print(f'Saved: {out}')
